import os
import docx
import openpyxl

os.makedirs("scratch", exist_ok=True)

# 1. Generate Mock DOCX Document
doc = docx.Document()
doc.add_heading("VinFast Warranty Policy", level=0)

doc.add_heading("Introduction", level=1)
doc.add_paragraph(
    "This document outlines the standard warranty guidelines for VinFast electric vehicles. "
    "Please read the sections below to understand the terms and duration of coverage."
)

doc.add_heading("Warranty Duration", level=1)
doc.add_paragraph(
    "The battery warranty for VF5 model covers a period of 7 years or 150,000 km, whichever comes first. "
    "For other commercial models, the standard warranty is 5 years."
)

doc.add_heading("Exclusions", level=1)
doc.add_paragraph(
    "Damage resulting from accidents, unauthorized modifications, or neglect is not covered under this warranty."
)

doc.save("scratch/sample_policy.docx")
print("Generated scratch/sample_policy.docx")

# 2. Generate Mock XLSX Document
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Models"

# Add Headers
ws.append(["Model", "Domain", "System", "Description"])

# Add Rows
ws.append(["VF5", "Aftersales", "Warranty Management", "Quy trình bảo hành pin xe VF5 kéo dài 7 năm"])
ws.append(["VF8", "Production", "Quality Control", "Kiểm định chất lượng lắp ráp khung gầm và hệ thống phanh"])
ws.append(["VF9", "Sales", "Pre-order System", "Hệ thống đặt cọc và ưu đãi khách hàng tiên phong"])

wb.save("scratch/sample_models.xlsx")
print("Generated scratch/sample_models.xlsx")
